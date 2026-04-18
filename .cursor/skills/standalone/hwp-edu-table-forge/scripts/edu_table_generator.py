#!/usr/bin/env python3
"""
Korean Education Table Generator

Generates high-quality DOCX tables for Korean elementary school teachers.
Supports 7 built-in template types with proper formatting per style-guide.md.

Usage:
    python edu_table_generator.py --data input.json --output output.docx
    python edu_table_generator.py --data input.json --output output.docx --landscape
    cat input.json | python edu_table_generator.py --output output.docx

Requirements:
    pip install python-docx
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

FONT_PRIMARY = "맑은 고딕"
FONT_FALLBACK = "Malgun Gothic"

COLOR_HEADER_BG = "F2F2F2"
COLOR_FIRST_COL_BG = "F9F9F9"
COLOR_WHITE = "FFFFFF"

BORDER_OUTER_SZ = 12        # 1.5pt outer frame (half-points)
BORDER_HEADER_BOTTOM_SZ = 8 # 1.0pt header bottom separator
BORDER_INNER_H_SZ = 4       # 0.5pt inner horizontal lines
BORDER_INNER_V_SZ = 2       # 0.25pt inner vertical lines
BORDER_INNER_H_COLOR = "808080"   # gray
BORDER_INNER_V_COLOR = "BFBFBF"   # light gray

CELL_MARGIN_TOP_TWIPS = 120    # 6pt
CELL_MARGIN_BOTTOM_TWIPS = 120 # 6pt
CELL_MARGIN_LEFT_TWIPS = 160   # 8pt
CELL_MARGIN_RIGHT_TWIPS = 160  # 8pt

A4_USABLE_MM = 170
A4_LANDSCAPE_USABLE_MM = 257


# ---------------------------------------------------------------------------
# Title helpers
# ---------------------------------------------------------------------------

def _ensure_suffix(value, suffix: str) -> str:
    """Append suffix only if value doesn't already contain it."""
    s = str(value).strip()
    if not s:
        return ""
    return s if suffix in s else f"{s}{suffix}"


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), attrs.get("val", "single"))
        el.set(qn("w:sz"), str(attrs.get("sz", BORDER_INNER_H_SZ)))
        el.set(qn("w:color"), attrs.get("color", "000000"))
        el.set(qn("w:space"), "0")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for edge, val in [("top", CELL_MARGIN_TOP_TWIPS),
                      ("bottom", CELL_MARGIN_BOTTOM_TWIPS),
                      ("start", CELL_MARGIN_LEFT_TWIPS),
                      ("end", CELL_MARGIN_RIGHT_TWIPS)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tcPr.append(margins)


def _set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_width(cell, width_mm: int):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_mm * 56.7)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_vertical_merge(cell, merge_type: str = "continue"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement("w:vMerge")
    if merge_type == "restart":
        vMerge.set(qn("w:val"), "restart")
    tcPr.append(vMerge)


def _set_row_height(row, height_mm: float):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_mm * 56.7)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def _remove_paragraph_spacing(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "60")
    spacing.set(qn("w:line"), "312")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


# ---------------------------------------------------------------------------
# Cell formatting helpers
# ---------------------------------------------------------------------------

def _format_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size_pt: float = 9,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    color: RGBColor | None = None,
):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for p in cell.paragraphs:
        p.clear()

    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = alignment
    _remove_paragraph_spacing(p)

    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = FONT_PRIMARY
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_PRIMARY)
        if i < len(lines) - 1:
            run.add_break()

    _set_cell_margins(cell)


def _format_header_cell(cell, text: str, size_pt: float = 11):
    _set_cell_shading(cell, COLOR_HEADER_BG)
    _format_cell_text(cell, text, bold=True, size_pt=size_pt,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_border(cell, bottom={
        "val": "single", "sz": BORDER_HEADER_BOTTOM_SZ,
        "color": BORDER_INNER_H_COLOR,
    })


def _apply_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if edge in ("top", "left", "bottom", "right"):
            el.set(qn("w:sz"), str(BORDER_OUTER_SZ))
            el.set(qn("w:color"), "000000")
        elif edge == "insideH":
            el.set(qn("w:sz"), str(BORDER_INNER_H_SZ))
            el.set(qn("w:color"), BORDER_INNER_H_COLOR)
        else:  # insideV
            el.set(qn("w:sz"), str(BORDER_INNER_V_SZ))
            el.set(qn("w:color"), BORDER_INNER_V_COLOR)
        el.set(qn("w:val"), "single")
        el.set(qn("w:space"), "0")
        borders.append(el)
    tblPr.append(borders)


def _set_table_width(table, width_mm: int):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_mm * 56.7)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def create_document(landscape: bool = False) -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_PRIMARY
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15
    rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
    if rFonts is not None:
        rFonts.set(qn("w:eastAsia"), FONT_PRIMARY)

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)

    return doc


def add_table_title(doc: Document, title: str, table_number: int | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    prefix = f"<표 {table_number}> " if table_number else ""
    run = p.add_run(f"{prefix}{title}")
    run.font.name = FONT_PRIMARY
    run.font.size = Pt(11)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_PRIMARY)
    p.paragraph_format.space_after = Pt(6)


def add_source_citation(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = FONT_PRIMARY
    run.font.size = Pt(8)
    run.italic = True
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_PRIMARY)


# ---------------------------------------------------------------------------
# Template 1: 교과 진도표 (Curriculum Progress)
# ---------------------------------------------------------------------------

def gen_curriculum_progress(doc: Document, data: dict, table_num: int = 1):
    sem = _ensure_suffix(data.get("semester", ""), "학기")
    grade = _ensure_suffix(data.get("grade", ""), "학년")
    title = f"{data.get('year', '')}학년도 {sem} {grade} {data.get('subject', '')} 교과 진도표"
    add_table_title(doc, title, table_num)

    headers = ["월", "주", "단원", "차시", "학습 내용", "교과서/비고"]
    col_widths = [15, 12, 35, 12, 55, 41]

    rows_data = []
    if data.get("months"):
        for month_block in data["months"]:
            month_num = str(month_block.get("month", ""))
            for wk in month_block.get("weeks", []):
                rows_data.append({
                    "month": month_num,
                    "week": str(wk.get("week", "")),
                    "unit": wk.get("unit", ""),
                    "periods": str(wk.get("hours", "")),
                    "content": wk.get("topic", ""),
                    "textbook": wk.get("note", ""),
                })
    else:
        for unit in data.get("units", []):
            weeks = unit.get("planned_weeks", [])
            ppw = max(1, unit.get("periods", len(weeks)) // max(1, len(weeks)))
            for i, week in enumerate(weeks):
                month = week.split("/")[0] if "/" in week else ""
                lc = unit.get("learning_content", [])
                content = lc[i] if i < len(lc) else ""
                pages = unit.get("textbook_pages", "") if i == 0 else ""
                rows_data.append({
                    "month": month,
                    "week": week,
                    "unit": f"{unit['unit_number']}. {unit['unit_title']}",
                    "periods": str(ppw),
                    "content": content,
                    "textbook": pages,
                })

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(table)
    _set_table_width(table, A4_USABLE_MM)

    for ci, header in enumerate(headers):
        _format_header_cell(table.rows[0].cells[ci], header)
        _set_cell_width(table.rows[0].cells[ci], col_widths[ci])

    prev_month = None
    prev_unit = None

    for ri, rd in enumerate(rows_data):
        row = table.rows[ri + 1]
        vals = [rd["month"], rd["week"], rd["unit"], rd["periods"],
                rd["content"], rd["textbook"]]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
        ]
        for ci, val in enumerate(vals):
            _format_cell_text(row.cells[ci], val, alignment=aligns[ci])
            _set_cell_width(row.cells[ci], col_widths[ci])

        _set_cell_shading(row.cells[0], COLOR_FIRST_COL_BG)

        cur_month = rd["month"]
        if cur_month == prev_month and prev_month:
            _set_cell_vertical_merge(row.cells[0], "continue")
        elif cur_month:
            _set_cell_vertical_merge(row.cells[0], "restart")
        prev_month = cur_month

        cur_unit = rd["unit"]
        if cur_unit == prev_unit and prev_unit:
            _set_cell_vertical_merge(row.cells[2], "continue")
        elif cur_unit:
            _set_cell_vertical_merge(row.cells[2], "restart")
        prev_unit = cur_unit


# ---------------------------------------------------------------------------
# Template 2: 교수학습과정안 (Lesson Plan)
# ---------------------------------------------------------------------------

def gen_lesson_plan(doc: Document, data: dict, table_num: int = 2):
    add_table_title(doc, "교수·학습 과정안", table_num)

    grade = _ensure_suffix(data.get("grade", ""), "학년")
    period_val = str(data.get("period", ""))
    period_display = period_val if "교시" in period_val else f"{period_val}교시"
    teacher = data.get("teacher_name", data.get("teacher", ""))
    unit_title = data.get("unit_title", data.get("unit", ""))
    lesson_num = data.get("lesson_number", "")
    total_lessons = data.get("total_lessons", "")
    lesson_display = f"{lesson_num}/{total_lessons}" if lesson_num else ""
    core_comp = data.get("core_competency", data.get("achievement_standard", ""))

    info_rows = [
        ("교 과", data.get("subject", ""), "학 년", grade),
        ("단 원", unit_title, "차 시", lesson_display),
        ("일 시", data.get("date", ""), "교시", period_display),
        ("핵심역량", core_comp, "지도교사", teacher),
    ]

    info_table = doc.add_table(rows=len(info_rows), cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(info_table)
    _set_table_width(info_table, A4_USABLE_MM)

    label_width = 25
    val_width = 60

    for ri, (l1, v1, l2, v2) in enumerate(info_rows):
        cells = info_table.rows[ri].cells
        _format_cell_text(cells[0], l1, bold=True, size_pt=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(cells[0], COLOR_HEADER_BG)
        _set_cell_width(cells[0], label_width)
        _format_cell_text(cells[1], v1, size_pt=10)
        _set_cell_width(cells[1], val_width)
        _format_cell_text(cells[2], l2, bold=True, size_pt=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(cells[2], COLOR_HEADER_BG)
        _set_cell_width(cells[2], label_width)
        _format_cell_text(cells[3], v2, size_pt=10)
        _set_cell_width(cells[3], val_width)

    objectives = data.get("lesson_objectives",
                          data.get("objectives",
                                   [data.get("learning_objective", "")]))
    obj_text = "\n".join(objectives) if isinstance(objectives, list) else str(objectives)
    obj_table = doc.add_table(rows=1, cols=1)
    obj_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(obj_table)
    _set_table_width(obj_table, A4_USABLE_MM)
    _format_cell_text(obj_table.rows[0].cells[0],
                      f"학습 목표: {obj_text}", bold=True, size_pt=10)

    flat_activities = _flatten_activities(data)
    act_headers = ["단계", "시간(분)", "교사 활동", "학생 활동", "자료 및 유의점(※)"]
    act_widths = [18, 15, 47, 47, 43]

    act_table = doc.add_table(rows=1 + len(flat_activities), cols=5)
    act_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(act_table)
    _set_table_width(act_table, A4_USABLE_MM)

    for ci, h in enumerate(act_headers):
        _format_header_cell(act_table.rows[0].cells[ci], h)
        _set_cell_width(act_table.rows[0].cells[ci], act_widths[ci])

    prev_phase = None
    for ri, act in enumerate(flat_activities):
        row = act_table.rows[ri + 1]
        phase = act["phase"]
        vals = [phase, act["duration"], act["teacher"], act["student"], act["materials"]]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
        ]
        for ci, val in enumerate(vals):
            _format_cell_text(row.cells[ci], val, alignment=aligns[ci])
            _set_cell_width(row.cells[ci], act_widths[ci])

        _set_cell_shading(row.cells[0], COLOR_FIRST_COL_BG)

        if phase == prev_phase and prev_phase:
            _set_cell_vertical_merge(row.cells[0], "continue")
        else:
            _set_cell_vertical_merge(row.cells[0], "restart")
        prev_phase = phase


def _flatten_activities(data: dict) -> list[dict]:
    """Normalize both canonical 'phases' format and flat 'activities' format."""
    flat = []
    if data.get("phases"):
        for phase_block in data["phases"]:
            phase_name = phase_block.get("phase", "")
            duration = str(phase_block.get("duration_min",
                                           phase_block.get("duration_minutes", "")))
            for act in phase_block.get("activities", []):
                flat.append({
                    "phase": phase_name,
                    "duration": duration if not flat or flat[-1]["phase"] != phase_name else "",
                    "teacher": act.get("activity", act.get("teacher_activity", "")),
                    "student": act.get("student_activity", ""),
                    "materials": act.get("materials", act.get("materials_notes", "")),
                })
    elif data.get("activities"):
        for act in data["activities"]:
            flat.append({
                "phase": act.get("phase", ""),
                "duration": str(act.get("duration_min",
                                        act.get("duration_minutes", ""))),
                "teacher": act.get("teacher_activity", ""),
                "student": act.get("student_activity", ""),
                "materials": act.get("materials_notes", ""),
            })
    return flat


# ---------------------------------------------------------------------------
# Template 3: 평가 루브릭 (Assessment Rubric)
# ---------------------------------------------------------------------------

def gen_rubric(doc: Document, data: dict, table_num: int = 3):
    assessment_title = data.get("assessment_title", "수행평가 루브릭")
    add_table_title(doc, f"{assessment_title} 루브릭", table_num)

    grade = _ensure_suffix(data.get("grade", ""), "학년")
    info_items = [
        ("교과", data.get("subject", "")),
        ("학년", grade),
        ("성취기준", data.get("achievement_standard", "")),
        ("평가유형", data.get("assessment_type", "")),
    ]
    info_table = doc.add_table(rows=len(info_items), cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(info_table)
    _set_table_width(info_table, A4_USABLE_MM)

    for ri, (label, value) in enumerate(info_items):
        _format_cell_text(info_table.rows[ri].cells[0], label, bold=True,
                          size_pt=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(info_table.rows[ri].cells[0], COLOR_HEADER_BG)
        _set_cell_width(info_table.rows[ri].cells[0], 30)
        _format_cell_text(info_table.rows[ri].cells[1], value, size_pt=10)
        _set_cell_width(info_table.rows[ri].cells[1], 140)

    criteria = data.get("criteria", [])
    level_names = _detect_level_names(criteria)
    num_levels = len(level_names)

    rubric_headers = ["평가기준", "비중(%)"] + level_names
    base_widths = [30, 15]
    remaining = A4_USABLE_MM - sum(base_widths)
    level_w = remaining // max(1, num_levels)
    rubric_widths = base_widths + [level_w] * num_levels

    rubric_table = doc.add_table(rows=1 + len(criteria), cols=len(rubric_headers))
    rubric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(rubric_table)
    _set_table_width(rubric_table, A4_USABLE_MM)

    for ci, h in enumerate(rubric_headers):
        _format_header_cell(rubric_table.rows[0].cells[ci], h)
        _set_cell_width(rubric_table.rows[0].cells[ci], rubric_widths[ci])

    for ri, crit in enumerate(criteria):
        row = rubric_table.rows[ri + 1]
        _format_cell_text(row.cells[0], crit.get("criterion", ""), bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(row.cells[0], COLOR_FIRST_COL_BG)
        _set_cell_width(row.cells[0], rubric_widths[0])

        weight = crit.get("weight_percent", crit.get("weight_pct", ""))
        _format_cell_text(row.cells[1], str(weight),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_width(row.cells[1], rubric_widths[1])

        levels = crit.get("levels", {})
        for li, level_name in enumerate(level_names):
            _format_cell_text(row.cells[2 + li], levels.get(level_name, ""))
            _set_cell_width(row.cells[2 + li], rubric_widths[2 + li])

    if data.get("achievement_standard"):
        add_source_citation(doc, "출처: 2022 개정 교육과정 성취기준 (교육부)")


def _detect_level_names(criteria: list[dict]) -> list[str]:
    """Auto-detect rubric level names from the data."""
    if not criteria:
        return ["상", "중", "하"]
    first_levels = criteria[0].get("levels", {})
    keys = list(first_levels.keys())
    if keys:
        return keys
    return ["상", "중", "하"]


# ---------------------------------------------------------------------------
# Template 4: 학생 관찰 기록표 (Student Observation Record)
# ---------------------------------------------------------------------------

def gen_observation_record(doc: Document, data: dict, table_num: int = 4):
    class_info = data.get("class_number", data.get("grade", ""))
    subject = data.get("subject", "")
    period = data.get("observation_period", "")
    subject_part = f" {subject}" if subject else ""
    title = f"{class_info}{subject_part} 학생 관찰 기록표 ({period})"
    add_table_title(doc, title, table_num)

    domains = data.get("domains", data.get("categories", []))
    fixed_cols = ["번호", "이름"]
    all_headers = fixed_cols + domains + ["종합 관찰 소견"]
    num_cols = len(all_headers)

    domain_width = min(18, max(12, (A4_USABLE_MM - 15 - 20 - 50) // max(1, len(domains))))
    col_widths = [15, 20] + [domain_width] * len(domains) + [
        A4_USABLE_MM - 15 - 20 - domain_width * len(domains)
    ]

    students = data.get("students", [])
    table = doc.add_table(rows=1 + len(students), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(table)
    _set_table_width(table, A4_USABLE_MM)

    for ci, h in enumerate(all_headers):
        _format_header_cell(table.rows[0].cells[ci], h)
        _set_cell_width(table.rows[0].cells[ci], col_widths[ci])

    for ri, student in enumerate(students):
        row = table.rows[ri + 1]
        _format_cell_text(row.cells[0], str(student.get("number", ri + 1)),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(row.cells[0], COLOR_FIRST_COL_BG)
        _set_cell_width(row.cells[0], col_widths[0])

        _format_cell_text(row.cells[1], student.get("name", ""),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_width(row.cells[1], col_widths[1])

        records = student.get("records", student.get("ratings", {}))
        for di, domain in enumerate(domains):
            _format_cell_text(row.cells[2 + di], records.get(domain, ""),
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_width(row.cells[2 + di], col_widths[2 + di])

        narrative_idx = 2 + len(domains)
        narrative = student.get("narrative", student.get("note", ""))
        _format_cell_text(row.cells[narrative_idx], narrative)
        _set_cell_width(row.cells[narrative_idx], col_widths[narrative_idx])


# ---------------------------------------------------------------------------
# Template 5: 연구 과제 보고서 표 (Research Report Table)
# ---------------------------------------------------------------------------

def gen_research_report(doc: Document, data: dict, table_num: int = 5):
    add_table_title(doc, "연구 과제 실행 계획", table_num)

    research_period = data.get("research_period", f"{data.get('year', '')}년")
    header_items = [
        ("연구 주제", data.get("title", "")),
        ("연구자", data.get("researcher", "")),
        ("소속", data.get("school", "")),
        ("연구 기간", research_period),
    ]
    header_items = [(k, v) for k, v in header_items if v]

    header_table = doc.add_table(rows=len(header_items), cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(header_table)
    _set_table_width(header_table, A4_USABLE_MM)

    for ri, (label, value) in enumerate(header_items):
        _format_cell_text(header_table.rows[ri].cells[0], label, bold=True,
                          size_pt=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(header_table.rows[ri].cells[0], COLOR_HEADER_BG)
        _set_cell_width(header_table.rows[ri].cells[0], 30)
        _format_cell_text(header_table.rows[ri].cells[1], value, size_pt=10)
        _set_cell_width(header_table.rows[ri].cells[1], 140)

    for section in data.get("sections", []):
        section_title = section.get("section_title", "")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(f"[{section_title}]")
        run.font.name = FONT_PRIMARY
        run.font.size = Pt(10)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_PRIMARY)

        rows = section.get("rows", [])
        if not rows:
            continue

        section_headers = section.get("headers", [])
        first_row = rows[0]

        if isinstance(first_row, list):
            num_cols = len(section_headers) if section_headers else len(first_row)
            col_w = A4_USABLE_MM // max(1, num_cols)
            tbl = doc.add_table(rows=1 + len(rows), cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _apply_table_borders(tbl)
            _set_table_width(tbl, A4_USABLE_MM)
            for ci, h in enumerate(section_headers[:num_cols]):
                _format_header_cell(tbl.rows[0].cells[ci], h)
                _set_cell_width(tbl.rows[0].cells[ci], col_w)
            for ri, row_data in enumerate(rows):
                for ci, val in enumerate(row_data[:num_cols]):
                    align = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    _format_cell_text(tbl.rows[ri + 1].cells[ci], str(val), alignment=align)
                    _set_cell_width(tbl.rows[ri + 1].cells[ci], col_w)
                _set_cell_shading(tbl.rows[ri + 1].cells[0], COLOR_FIRST_COL_BG)

        elif isinstance(first_row, dict):
            if "category" in first_row:
                tbl = doc.add_table(rows=len(rows), cols=2)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                _apply_table_borders(tbl)
                _set_table_width(tbl, A4_USABLE_MM)
                for ri, row_data in enumerate(rows):
                    _format_cell_text(tbl.rows[ri].cells[0], row_data.get("category", ""),
                                      bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    _set_cell_shading(tbl.rows[ri].cells[0], COLOR_HEADER_BG)
                    _set_cell_width(tbl.rows[ri].cells[0], 30)
                    _format_cell_text(tbl.rows[ri].cells[1], row_data.get("content", ""))
                    _set_cell_width(tbl.rows[ri].cells[1], 140)
            elif "phase" in first_row:
                headers = ["단계", "기간", "과제 내용", "산출물"]
                widths = [20, 20, 75, 55]
                tbl = doc.add_table(rows=1 + len(rows), cols=4)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                _apply_table_borders(tbl)
                _set_table_width(tbl, A4_USABLE_MM)
                for ci, h in enumerate(headers):
                    _format_header_cell(tbl.rows[0].cells[ci], h)
                    _set_cell_width(tbl.rows[0].cells[ci], widths[ci])
                for ri, row_data in enumerate(rows):
                    r = tbl.rows[ri + 1]
                    vals = [row_data.get("phase", ""), row_data.get("period", ""),
                            row_data.get("task", ""), row_data.get("output", "")]
                    aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                              WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
                    for ci, val in enumerate(vals):
                        _format_cell_text(r.cells[ci], val, alignment=aligns[ci])
                        _set_cell_width(r.cells[ci], widths[ci])
                    _set_cell_shading(r.cells[0], COLOR_FIRST_COL_BG)


# ---------------------------------------------------------------------------
# Template 6: 학급 운영 계획 (Classroom Management Plan)
# ---------------------------------------------------------------------------

def gen_classroom_management(doc: Document, data: dict, table_num: int = 6):
    grade_raw = str(data.get("grade", ""))
    year = data.get("year", "")
    title = f"{year}학년도 {grade_raw} 학급 운영 계획"
    add_table_title(doc, title, table_num)

    class_goal = data.get("class_goal", data.get("class_motto", ""))
    goal_table = doc.add_table(rows=1, cols=1)
    goal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(goal_table)
    _set_table_width(goal_table, A4_USABLE_MM)
    _format_cell_text(goal_table.rows[0].cells[0],
                      f"학급 목표: {class_goal}", bold=True, size_pt=10)

    plans = data.get("monthly_plans", data.get("monthly_plan", []))

    headers = ["월", "중점 지도/주제", "학급 활동"]
    widths = [15, 50, 105]

    table = doc.add_table(rows=1 + len(plans), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(table)
    _set_table_width(table, A4_USABLE_MM)

    for ci, h in enumerate(headers):
        _format_header_cell(table.rows[0].cells[ci], h)
        _set_cell_width(table.rows[0].cells[ci], widths[ci])

    for ri, plan in enumerate(plans):
        row = table.rows[ri + 1]
        month_val = str(plan.get("month", ""))
        focus = plan.get("focus", plan.get("theme", ""))
        activities_raw = plan.get("activities", plan.get("events", ""))
        if isinstance(activities_raw, list):
            activities_str = ", ".join(str(a) for a in activities_raw)
        else:
            activities_str = str(activities_raw)

        vals = [month_val, focus, activities_str]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.LEFT]
        for ci, val in enumerate(vals):
            _format_cell_text(row.cells[ci], val, alignment=aligns[ci])
            _set_cell_width(row.cells[ci], widths[ci])

        _set_cell_shading(row.cells[0], COLOR_FIRST_COL_BG)


# ---------------------------------------------------------------------------
# Template 7: 시간표 (Weekly Timetable)
# ---------------------------------------------------------------------------

def gen_timetable(doc: Document, data: dict, table_num: int = 7):
    grade_raw = str(data.get("grade", ""))
    sem = _ensure_suffix(data.get("semester", ""), "학기")
    title = f"{data.get('year', '')}학년도 {sem} {grade_raw} 주간 시간표"
    add_table_title(doc, title, table_num)

    day_names = ["월", "화", "수", "목", "금"]
    schedule = data.get("schedule", data.get("days", {}))

    periods_raw = data.get("periods", [])
    period_times_raw = data.get("period_times", [])

    actual_periods = []
    lunch_after = 0

    if periods_raw and isinstance(periods_raw[0], dict):
        for p in periods_raw:
            pnum = p.get("period", "")
            if pnum == "점심" or str(pnum).lower() == "lunch":
                lunch_after = len(actual_periods)
                continue
            time_str = f"{p.get('start', '')}-{p.get('end', '')}"
            actual_periods.append({"num": pnum, "time": time_str})
    elif period_times_raw:
        actual_periods = [{"num": i + 1, "time": t} for i, t in enumerate(period_times_raw)]
        lunch_after = data.get("lunch_after_period", 4)
    else:
        ppd = data.get("periods_per_day", 6)
        actual_periods = [{"num": i + 1, "time": ""} for i in range(ppd)]
        lunch_after = data.get("lunch_after_period", 4)

    periods_per_day = len(actual_periods)
    has_lunch = lunch_after > 0
    total_rows = 1 + periods_per_day + (1 if has_lunch else 0)

    headers = ["교시", "시간"] + day_names
    col_widths = [15, 30] + [(A4_USABLE_MM - 15 - 30) // 5] * 5

    table = doc.add_table(rows=total_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(table)
    _set_table_width(table, A4_USABLE_MM)

    for ci, h in enumerate(headers):
        _format_header_cell(table.rows[0].cells[ci], h)
        _set_cell_width(table.rows[0].cells[ci], col_widths[ci])

    display_row = 1
    period_data_idx = 0

    for period_data_idx in range(periods_per_day):
        if has_lunch and period_data_idx == lunch_after:
            lunch_row = table.rows[display_row]
            _format_cell_text(lunch_row.cells[0], "",
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_width(lunch_row.cells[0], col_widths[0])

            lunch_time = ""
            if lunch_after > 0 and lunch_after <= len(actual_periods):
                prev_end = actual_periods[lunch_after - 1]["time"].split("-")[-1]
                if lunch_after < len(actual_periods):
                    next_start = actual_periods[lunch_after]["time"].split("-")[0]
                    if prev_end and next_start:
                        lunch_time = f"{prev_end}-{next_start}"
            _format_cell_text(lunch_row.cells[1], lunch_time,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_width(lunch_row.cells[1], col_widths[1])

            merged_cell = lunch_row.cells[2]
            for merge_ci in range(3, len(headers)):
                merged_cell.merge(lunch_row.cells[merge_ci])
            _format_cell_text(merged_cell, "점심", bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10)
            _set_cell_shading(merged_cell, COLOR_FIRST_COL_BG)
            display_row += 1

        row = table.rows[display_row]
        p_info = actual_periods[period_data_idx]
        _format_cell_text(row.cells[0], str(p_info["num"]),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_width(row.cells[0], col_widths[0])

        _format_cell_text(row.cells[1], p_info["time"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_width(row.cells[1], col_widths[1])

        for di, day in enumerate(day_names):
            day_schedule = schedule.get(day, [])
            raw_subj = day_schedule[period_data_idx] if period_data_idx < len(day_schedule) else None
            subject = str(raw_subj) if raw_subj is not None else "—"
            _format_cell_text(row.cells[2 + di], subject,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_width(row.cells[2 + di], col_widths[2 + di])

        display_row += 1


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

GENERATORS = {
    "curriculum_progress": gen_curriculum_progress,
    "lesson_plan": gen_lesson_plan,
    "rubric": gen_rubric,
    "observation_record": gen_observation_record,
    "research_report": gen_research_report,
    "classroom_management": gen_classroom_management,
    "timetable": gen_timetable,
}


def generate_table(data: dict, output_path: str, landscape: bool = False,
                   table_number: int = 1) -> str:
    template_type = data.get("template_type", "")
    gen_fn = GENERATORS.get(template_type)
    if not gen_fn:
        supported = ", ".join(sorted(GENERATORS.keys()))
        raise ValueError(
            f"Unknown template_type '{template_type}'. Supported: {supported}"
        )

    doc = create_document(landscape=landscape)
    gen_fn(doc, data, table_num=table_number)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out.resolve())


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Korean Education Table Generator — DOCX output"
    )
    parser.add_argument("--data", type=str, help="Path to JSON input file (or use stdin)")
    parser.add_argument("--output", "-o", type=str, required=True, help="Output .docx path")
    parser.add_argument("--landscape", action="store_true", help="Use landscape orientation")
    parser.add_argument("--table-number", type=int, default=1, help="Table number prefix")
    args = parser.parse_args()

    if args.data:
        with open(args.data, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.load(sys.stdin)

    out_path = generate_table(data, args.output, landscape=args.landscape,
                              table_number=args.table_number)
    print(f"Generated: {out_path}")


if __name__ == "__main__":
    main()
